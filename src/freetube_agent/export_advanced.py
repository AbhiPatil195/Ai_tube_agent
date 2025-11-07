"""
Advanced Export Module
Exports transcripts and summaries to multiple formats: PDF, Word, Markdown, JSON
"""

from __future__ import annotations
from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime
import json

from .transcribe import Transcript
from .paths import TRANSCRIPTS


def export_to_markdown(
    transcript: Transcript,
    video_title: str,
    summary_data: Optional[Dict[str, Any]] = None,
    output_path: Optional[Path] = None
) -> Path:
    """
    Export transcript and optional summary to Markdown format.
    
    Args:
        transcript: Transcript object
        video_title: Title of the video
        summary_data: Optional summary data from summarize module
        output_path: Optional custom output path
    
    Returns:
        Path to exported markdown file
    """
    if output_path is None:
        output_path = TRANSCRIPTS / f"{video_title}_export.md"
    
    lines = []
    
    # Header
    lines.append(f"# {video_title}\n")
    lines.append(f"*Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n")
    lines.append("---\n")
    
    # Summary section if available
    if summary_data:
        lines.append("## 📝 Summary\n")
        
        if summary_data.get("tldr", {}).get("success"):
            lines.append("### TL;DR\n")
            lines.append(f"> {summary_data['tldr']['tldr']}\n")
        
        if summary_data.get("key_points", {}).get("success"):
            lines.append("### Key Points\n")
            for point in summary_data["key_points"]["key_points"]:
                lines.append(f"- {point}\n")
            lines.append("\n")
        
        if summary_data.get("topics", {}).get("success"):
            lines.append("### Topics Covered\n")
            for topic in summary_data["topics"]["topics"]:
                lines.append(f"**{topic['name']}**: {topic['description']}\n")
            lines.append("\n")
        
        if summary_data.get("summary", {}).get("success"):
            lines.append("### Detailed Summary\n")
            lines.append(f"{summary_data['summary']['summary']}\n\n")
        
        lines.append("---\n")
    
    # Full Transcript
    lines.append("## 📄 Full Transcript\n")
    
    if transcript.segments:
        for segment in transcript.segments:
            timestamp = format_timestamp(segment.start)
            lines.append(f"**[{timestamp}]** {segment.text}\n")
    else:
        lines.append(transcript.text)
    
    # Write to file
    output_path.write_text("\n".join(lines), encoding="utf-8")
    return output_path


def export_to_pdf(
    transcript: Transcript,
    video_title: str,
    summary_data: Optional[Dict[str, Any]] = None,
    output_path: Optional[Path] = None
) -> Path:
    """
    Export transcript and optional summary to PDF format.
    
    Args:
        transcript: Transcript object
        video_title: Title of the video
        summary_data: Optional summary data
        output_path: Optional custom output path
    
    Returns:
        Path to exported PDF file
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
    except ImportError:
        raise RuntimeError("reportlab not installed. Run: pip install reportlab")
    
    if output_path is None:
        output_path = TRANSCRIPTS / f"{video_title}_export.pdf"
    
    # Create PDF document
    doc = SimpleDocTemplate(str(output_path), pagesize=letter)
    story = []
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor='#FF0000',
        spaceAfter=12,
        alignment=TA_CENTER
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor='#333333',
        spaceAfter=10,
        spaceBefore=12
    )
    
    # Title
    story.append(Paragraph(video_title, title_style))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
    story.append(Spacer(1, 0.3*inch))
    
    # Summary section if available
    if summary_data:
        story.append(Paragraph("Summary", heading_style))
        
        if summary_data.get("tldr", {}).get("success"):
            story.append(Paragraph("<b>TL;DR:</b>", styles['Heading3']))
            story.append(Paragraph(summary_data['tldr']['tldr'], styles['Italic']))
            story.append(Spacer(1, 0.2*inch))
        
        if summary_data.get("key_points", {}).get("success"):
            story.append(Paragraph("<b>Key Points:</b>", styles['Heading3']))
            for point in summary_data["key_points"]["key_points"]:
                story.append(Paragraph(f"• {point}", styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
        
        if summary_data.get("topics", {}).get("success"):
            story.append(Paragraph("<b>Topics Covered:</b>", styles['Heading3']))
            for topic in summary_data["topics"]["topics"]:
                story.append(Paragraph(f"<b>{topic['name']}:</b> {topic['description']}", styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
        
        story.append(PageBreak())
    
    # Full Transcript
    story.append(Paragraph("Full Transcript", heading_style))
    story.append(Spacer(1, 0.1*inch))
    
    if transcript.segments:
        for segment in transcript.segments:
            timestamp = format_timestamp(segment.start)
            text = f"<b>[{timestamp}]</b> {segment.text}"
            story.append(Paragraph(text, styles['Normal']))
            story.append(Spacer(1, 0.05*inch))
    else:
        story.append(Paragraph(transcript.text, styles['Normal']))
    
    # Build PDF
    doc.build(story)
    return output_path


def export_to_word(
    transcript: Transcript,
    video_title: str,
    summary_data: Optional[Dict[str, Any]] = None,
    output_path: Optional[Path] = None
) -> Path:
    """
    Export transcript and optional summary to Word DOCX format.
    
    Args:
        transcript: Transcript object
        video_title: Title of the video
        summary_data: Optional summary data
        output_path: Optional custom output path
    
    Returns:
        Path to exported DOCX file
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")
    
    if output_path is None:
        output_path = TRANSCRIPTS / f"{video_title}_export.docx"
    
    doc = Document()
    
    # Title
    title = doc.add_heading(video_title, 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(255, 0, 0)  # YouTube red
    
    # Metadata
    meta = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta_run = meta.runs[0]
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()
    
    # Summary section if available
    if summary_data:
        doc.add_heading("Summary", 1)
        
        if summary_data.get("tldr", {}).get("success"):
            doc.add_heading("TL;DR", 2)
            tldr_para = doc.add_paragraph(summary_data['tldr']['tldr'])
            tldr_para.runs[0].font.italic = True
            doc.add_paragraph()
        
        if summary_data.get("key_points", {}).get("success"):
            doc.add_heading("Key Points", 2)
            for point in summary_data["key_points"]["key_points"]:
                doc.add_paragraph(point, style='List Bullet')
            doc.add_paragraph()
        
        if summary_data.get("topics", {}).get("success"):
            doc.add_heading("Topics Covered", 2)
            for topic in summary_data["topics"]["topics"]:
                p = doc.add_paragraph()
                p.add_run(f"{topic['name']}: ").bold = True
                p.add_run(topic['description'])
            doc.add_paragraph()
        
        if summary_data.get("summary", {}).get("success"):
            doc.add_heading("Detailed Summary", 2)
            doc.add_paragraph(summary_data['summary']['summary'])
        
        doc.add_page_break()
    
    # Full Transcript
    doc.add_heading("Full Transcript", 1)
    
    if transcript.segments:
        for segment in transcript.segments:
            timestamp = format_timestamp(segment.start)
            p = doc.add_paragraph()
            timestamp_run = p.add_run(f"[{timestamp}] ")
            timestamp_run.bold = True
            timestamp_run.font.color.rgb = RGBColor(62, 166, 255)  # YouTube blue
            p.add_run(segment.text)
    else:
        doc.add_paragraph(transcript.text)
    
    # Save
    doc.save(str(output_path))
    return output_path


def export_to_json(
    transcript: Transcript,
    video_title: str,
    summary_data: Optional[Dict[str, Any]] = None,
    output_path: Optional[Path] = None
) -> Path:
    """
    Export transcript and optional summary to JSON format.
    
    Args:
        transcript: Transcript object
        video_title: Title of the video
        summary_data: Optional summary data
        output_path: Optional custom output path
    
    Returns:
        Path to exported JSON file
    """
    if output_path is None:
        output_path = TRANSCRIPTS / f"{video_title}_export.json"
    
    # Build JSON structure
    data = {
        "title": video_title,
        "generated_at": datetime.now().isoformat(),
        "transcript": {
            "text": transcript.text,
            "segments": [
                {
                    "start": seg.start,
                    "end": seg.end,
                    "duration": seg.end - seg.start,
                    "text": seg.text
                }
                for seg in transcript.segments
            ] if transcript.segments else [],
            "word_count": len(transcript.text.split()),
            "duration": transcript.segments[-1].end if transcript.segments else 0
        }
    }
    
    # Add summary if available
    if summary_data:
        data["summary"] = summary_data
    
    # Write JSON
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    
    return output_path


def export_blog_post(
    transcript: Transcript,
    video_title: str,
    summary_data: Optional[Dict[str, Any]] = None,
    output_path: Optional[Path] = None,
    include_metadata: bool = True
) -> Path:
    """
    Export as a blog post format with HTML.
    
    Args:
        transcript: Transcript object
        video_title: Title of the video
        summary_data: Optional summary data
        output_path: Optional custom output path
        include_metadata: Include SEO metadata
    
    Returns:
        Path to exported HTML file
    """
    if output_path is None:
        output_path = TRANSCRIPTS / f"{video_title}_blog.html"
    
    html = []
    
    if include_metadata:
        html.append("<!DOCTYPE html>")
        html.append("<html lang='en'>")
        html.append("<head>")
        html.append(f"<title>{video_title}</title>")
        html.append("<meta charset='UTF-8'>")
        html.append("<meta name='viewport' content='width=device-width, initial-scale=1.0'>")
        html.append("<style>")
        html.append("body { font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.6; }")
        html.append("h1 { color: #FF0000; }")
        html.append("h2 { color: #333; border-bottom: 2px solid #FF0000; padding-bottom: 10px; }")
        html.append(".tldr { background: #f5f5f5; padding: 15px; border-left: 4px solid #FF0000; margin: 20px 0; }")
        html.append(".timestamp { color: #3ea6ff; font-weight: bold; }")
        html.append(".key-points { background: #fafafa; padding: 20px; border-radius: 8px; }")
        html.append("ul { padding-left: 20px; }")
        html.append("</style>")
        html.append("</head>")
        html.append("<body>")
    
    # Title
    html.append(f"<h1>{video_title}</h1>")
    html.append(f"<p><em>Published: {datetime.now().strftime('%B %d, %Y')}</em></p>")
    
    # Summary section
    if summary_data:
        if summary_data.get("tldr", {}).get("success"):
            html.append("<div class='tldr'>")
            html.append(f"<strong>TL;DR:</strong> {summary_data['tldr']['tldr']}")
            html.append("</div>")
        
        if summary_data.get("key_points", {}).get("success"):
            html.append("<div class='key-points'>")
            html.append("<h2>Key Takeaways</h2>")
            html.append("<ul>")
            for point in summary_data["key_points"]["key_points"]:
                html.append(f"<li>{point}</li>")
            html.append("</ul>")
            html.append("</div>")
        
        if summary_data.get("summary", {}).get("success"):
            html.append("<h2>Overview</h2>")
            html.append(f"<p>{summary_data['summary']['summary']}</p>")
    
    # Transcript
    html.append("<h2>Full Transcript</h2>")
    
    if transcript.segments:
        for segment in transcript.segments:
            timestamp = format_timestamp(segment.start)
            html.append(f"<p><span class='timestamp'>[{timestamp}]</span> {segment.text}</p>")
    else:
        html.append(f"<p>{transcript.text}</p>")
    
    if include_metadata:
        html.append("</body>")
        html.append("</html>")
    
    # Write HTML
    output_path.write_text("\n".join(html), encoding="utf-8")
    return output_path


def format_timestamp(seconds: float) -> str:
    """Format seconds as MM:SS or HH:MM:SS"""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    
    if hours > 0:
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    return f"{minutes:02d}:{secs:02d}"
